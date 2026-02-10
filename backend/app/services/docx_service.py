import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

def add_markdown_content(doc, text: str):
    """
    Parses simple markdown text and adds it to the docx Document.
    """
    if not text:
        return

    try:
        lines = text.split('\n')
        in_code_block = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Code blocks
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
                
            if in_code_block:
                try:
                    p = doc.add_paragraph(line)
                    # p.style = 'Quote' # Disable style for safety if template missing
                    p.paragraph_format.left_indent = Pt(20)
                except Exception as e:
                    logger.warning(f"Error adding code block line: {e}")
                    doc.add_paragraph(line)
                continue

            # Headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                _add_formatted_run(p, line[2:])
            # Normal paragraph
            else:
                p = doc.add_paragraph()
                _add_formatted_run(p, line)
    except Exception as e:
        logger.error(f"Error in add_markdown_content: {e}")
        doc.add_paragraph(text) # Fallback to plain text

def _add_formatted_run(paragraph, text):
    """
    Helper to process inline formatting (bold).
    """
    try:
        # Split by bold markers **
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)
    except Exception as e:
        logger.error(f"Error in _add_formatted_run: {e}")
        paragraph.add_run(text)

def _set_default_font(doc):
    """Set default font for Normal and Headings to Microsoft YaHei."""
    # Styles to update
    styles = ['Normal'] + [f'Heading {i}' for i in range(1, 10)]
    
    for style_name in styles:
        try:
            if style_name in doc.styles:
                style = doc.styles[style_name]
                style.font.name = 'Microsoft YaHei'
                style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        except Exception:
            pass

def generate_report_docx(report_data: dict) -> BytesIO:
    try:
        doc = Document()
        _set_default_font(doc)
        
        # Title
        username = report_data.get('username', 'Unknown')
        # Handle both camelCase and snake_case
        date_val = report_data.get('created_at') or report_data.get('createdAt')
        date_str = _format_date(date_val)
            
        doc.add_heading(f"周报 - {username} ({date_str})", 0)
        
        # Metadata
        p = doc.add_paragraph()
        p.add_run(f"提交人: {username}\n").bold = True
        p.add_run(f"提交时间: {date_str}\n")
        
        # Content
        doc.add_heading("本周工作内容", level=1)
        add_markdown_content(doc, report_data.get('content', '无内容'))
        
        # Details
        if report_data.get('details'):
            doc.add_heading("详细项目汇报", level=1)
            for detail in report_data['details']:
                doc.add_heading(detail.get('project_title', 'Unknown Project'), level=2)
                if detail.get('content'):
                    p = doc.add_paragraph()
                    p.add_run("进展: ").bold = True
                    add_markdown_content(doc, detail['content'])
                if detail.get('plan'):
                    p = doc.add_paragraph()
                    p.add_run("计划: ").bold = True
                    add_markdown_content(doc, detail['plan'])

        # Footer
        _add_footer(doc)

        return _save_doc(doc)
    except Exception as e:
        logger.error(f"Error generating report docx: {e}")
        raise

def generate_timeline_docx(project_title: str, events: list) -> BytesIO:
    try:
        doc = Document()
        _set_default_font(doc)
        
        doc.add_heading(f"项目时间轴 - {project_title}", 0)
        doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        
        for event in events:
            _add_event_section(doc, event)
            doc.add_paragraph("-" * 20) 

        return _save_doc(doc)
    except Exception as e:
        logger.error(f"Error generating timeline docx: {e}")
        raise

def generate_event_docx(event: dict) -> BytesIO:
    """
    Generate a Word document for a single event.
    """
    try:
        logger.info(f"Generating event docx for event: {event.get('id')}")
        doc = Document()
        _set_default_font(doc)
        
        # Only export content for single event
        content = event.get('content', '')
        add_markdown_content(doc, content)
        
        return _save_doc(doc)
    except Exception as e:
        logger.error(f"Error generating event docx: {e}")
        raise

def _add_event_section(doc, event):
    try:
        date_val = event.get('date')
        date_str = _format_date(date_val)
        
        event_type = event.get('type')
        if hasattr(event_type, 'value'):
            event_type = event_type.value
        elif not event_type:
            event_type = 'EVENT'
            
        author = event.get('author_name', 'Unknown')
        
        heading = doc.add_heading(f"{date_str} - {event_type} ({author})", level=2)
        
        content = event.get('content', '')
        add_markdown_content(doc, content)
        
        attachments = event.get('attachments', [])
        if attachments:
            p = doc.add_paragraph()
            p.add_run("附件: ").bold = True
            att_names = [att.get('name', 'file') for att in attachments]
            p.add_run(", ".join(att_names))
    except Exception as e:
        logger.error(f"Error in _add_event_section: {e}")
        doc.add_paragraph(f"[Error rendering event section: {e}]")

def _format_date(date_val):
    if not date_val: return ''
    try:
        if isinstance(date_val, str):
            date_obj = datetime.fromisoformat(date_val.replace('Z', '+00:00'))
            return date_obj.strftime('%Y-%m-%d %H:%M')
        elif isinstance(date_val, datetime):
            return date_val.strftime('%Y-%m-%d %H:%M')
    except ValueError:
        pass
    return str(date_val)

def _add_footer(doc):
    try:
        section = doc.sections[0]
        footer = section.footer
        footer.paragraphs[0].text = f"Generated by DeptSync at {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    except:
        pass

def _save_doc(doc):
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
